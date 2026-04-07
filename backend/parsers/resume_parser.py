"""
Resume Parser Module
Extracts raw text from PDF, DOCX, and TXT resume files.
"""

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using pdfplumber."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                # Try to extract tables as structured text
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        if row:
                            text_parts.append("  |  ".join(
                                cell.strip() for cell in row if cell
                            ))

                # Extract main text
                page_text = page.extract_text(x_tolerance=3, y_tolerance=3)
                if page_text:
                    text_parts.append(page_text)

        parsed_text = '\n'.join(text_parts).strip()
        
        # OCR Fallback
        if len(parsed_text) < 50:
            logger.info("PDF text too short, falling back to OCR...")
            try:
                from pdf2image import convert_from_bytes
                import pytesseract
                images = convert_from_bytes(file_bytes)
                ocr_parts = []
                for img in images:
                    ocr_parts.append(pytesseract.image_to_string(img))
                parsed_text = '\n'.join(ocr_parts).strip()
            except Exception as ocr_err:
                logger.warning(f"OCR fallback failed: {ocr_err}")

        return parsed_text
    except ImportError:
        raise RuntimeError("pdfplumber not installed. Run: pip install pdfplumber")
    except Exception as e:
        logger.error(f"PDF parsing error: {e}")
        raise RuntimeError(f"Could not parse PDF: {e}")


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = "  |  ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        return '\n'.join(text_parts)
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        logger.error(f"DOCX parsing error: {e}")
        raise RuntimeError(f"Could not parse DOCX: {e}")


def parse_txt(file_bytes: bytes) -> str:
    """Decode plain text files, trying common encodings."""
    for encoding in ('utf-8', 'latin-1', 'windows-1252', 'ascii'):
        try:
            return file_bytes.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    raise RuntimeError("Could not decode text file — unsupported encoding.")


def parse_resume_file(filename: str, file_bytes: bytes) -> str:
    """
    Dispatch to the correct parser based on file extension.
    Returns raw text extracted from the resume.
    """
    ext = Path(filename).suffix.lower()

    if ext == '.pdf':
        raw_text = parse_pdf(file_bytes)
    elif ext in ('.docx', '.doc'):
        raw_text = parse_docx(file_bytes)
    elif ext in ('.txt', '.text', '.md'):
        raw_text = parse_txt(file_bytes)
    else:
        raise ValueError(
            f"Unsupported file format: '{ext}'. "
            "Please upload a PDF, DOCX, or TXT file."
        )

    # Basic cleanup
    raw_text = _clean_text(raw_text)

    if len(raw_text.strip()) < 50:
        raise ValueError("Extracted text is too short — the file may be empty or image-only.")

    return raw_text


def _clean_text(text: str) -> str:
    """
    Clean up extracted text: normalize whitespace, remove junk characters,
    collapse repeated blank lines.
    """
    import re

    # Replace non-breaking spaces, form feeds, carriage returns
    text = text.replace('\xa0', ' ').replace('\r', '\n').replace('\f', '\n')

    # Remove control characters except newlines
    text = re.sub(r'[^\S\n]+', ' ', text)           # multi-space → single space
    text = re.sub(r'\n{3,}', '\n\n', text)           # 3+ blank lines → 2
    text = re.sub(r'[ \t]+\n', '\n', text)           # trailing spaces
    text = re.sub(r'\n[ \t]+', '\n', text)           # leading spaces on lines

    return text.strip()
